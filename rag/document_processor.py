"""
Módulo para procesamiento de documentos e imágenes
"""
import os
import re
import hashlib
from typing import List, Set, Optional
from PyPDF2 import PdfReader
import docx
from models import Document, DocumentTuple
from config import KNOWLEDGE_DIR

# Importar el servicio de visión si está disponible
try:
    from services.vision_service import VisionService
    VISION_AVAILABLE = True
except ImportError:
    VISION_AVAILABLE = False
    print("⚠️ Servicio de visión no disponible. Instale las dependencias para procesamiento de imágenes.")


class DocumentProcessor:
    """Procesador de documentos para el sistema RAG"""
    
    # Extensiones de archivo soportadas
    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}
    # Extensiones de imagen soportadas
    IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"}
    # Archivos a ignorar explícitamente
    IGNORED_FILENAMES = {"hashes.txt"}
    
    def __init__(self):
        """Inicializar el procesador de documentos"""
        self.vision_service = None
        if VISION_AVAILABLE:
            try:
                self.vision_service = VisionService()
            except Exception as e:
                print(f"⚠️ Error inicializando servicio de visión: {e}")
                self.vision_service = None
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Limpia y normaliza el texto extraído
        
        Args:
            text: Texto a limpiar
            
        Returns:
            Texto limpio
        """
        # Remover caracteres especiales no imprimibles
        text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\xA1-\xFF]", "", text)
        # Remover líneas vacías y espacios extras
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
    
    def extract_text_from_pdf(self, path: str) -> str:
        """
        Extrae texto de un archivo PDF
        
        Args:
            path: Ruta al archivo PDF
            
        Returns:
            Texto extraído del PDF
        """
        try:
            reader = PdfReader(path)
            text = "\n".join(
                page.extract_text() 
                for page in reader.pages 
                if page.extract_text()
            )
            return DocumentProcessor.clean_text(text)
        except Exception as e:
            print(f"❌ Error leyendo PDF {path}: {e}")
            return ""
    
    def extract_text_from_docx(self, path: str) -> str:
        """
        Extrae texto de un archivo DOCX
        
        Args:
            path: Ruta al archivo DOCX
            
        Returns:
            Texto extraído del DOCX
        """
        try:
            doc = docx.Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return DocumentProcessor.clean_text(text)
        except Exception as e:
            print(f"❌ Error leyendo DOCX {path}: {e}")
            return ""
    
    def load_text_file(self, filepath: str) -> str:
        """
        Carga contenido de un archivo de texto
        
        Args:
            filepath: Ruta al archivo de texto
            
        Returns:
            Contenido del archivo
        """
        try:
            with open(filepath, encoding="utf-8") as f:
                content = f.read()
            return DocumentProcessor.clean_text(content)
        except Exception as e:
            print(f"❌ Error leyendo archivo {filepath}: {e}")
            return ""
    
    def extract_text_from_file(self, filepath: str) -> str:
        """
        Extrae texto de un archivo según su extensión
        
        Args:
            filepath: Ruta al archivo
            
        Returns:
            Texto extraído del archivo o descripción de imagen
        """
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext == ".pdf":
            return self.extract_text_from_pdf(filepath)
        elif ext == ".docx":
            return self.extract_text_from_docx(filepath)
        elif ext == ".txt":
            return self.load_text_file(filepath)
        elif ext in self.IMAGE_EXTENSIONS:
            return self.extract_description_from_image(filepath)
        else:
            print(f"⚠️ Tipo de archivo no soportado: {ext}")
            return ""
    
    def extract_description_from_image(self, filepath: str) -> str:
        """
        Extrae descripción cultural de una imagen
        
        Args:
            filepath: Ruta al archivo de imagen
            
        Returns:
            Descripción textual de la imagen para indexación
        """
        if not self.vision_service:
            return f"Imagen: {os.path.basename(filepath)} - Análisis visual no disponible"
        
        try:
            analysis = self.vision_service.analyze_image(filepath)
            
            if "error" in analysis:
                return f"Imagen: {os.path.basename(filepath)} - Error en análisis: {analysis['error']}"
            
            # Construir descripción textual para indexación
            description_parts = []
            description_parts.append(f"IMAGEN: {os.path.basename(filepath)}")
            
            if analysis.get("description"):
                description_parts.append(f"Descripción general: {analysis['description']}")
            
            if analysis.get("cultural_objects"):
                for obj in analysis["cultural_objects"]:
                    description_parts.append(f"Objeto cultural identificado: {obj['name']} de la cultura {obj['culture']}")
                    description_parts.append(f"Significado: {obj['significance']}")
                    description_parts.append(f"Descripción: {obj['description']}")
            
            if analysis.get("analysis_summary"):
                description_parts.append(f"Análisis cultural: {analysis['analysis_summary']}")
            
            # Agregar información de colores dominantes si están disponibles
            if analysis.get("dominant_colors"):
                color_info = "Colores dominantes presentes en la imagen"
                description_parts.append(color_info)
            
            return "\n".join(description_parts)
            
        except Exception as e:
            print(f"⚠️ Error analizando imagen {filepath}: {e}")
            return f"Imagen: {os.path.basename(filepath)} - Error en procesamiento: {str(e)}"
    
    def is_supported_file(self, filepath: str) -> bool:
        """
        Verifica si el archivo es de un tipo soportado (documentos o imágenes)
        
        Args:
            filepath: Ruta al archivo
            
        Returns:
            True si el archivo es soportado
        """
        filename = os.path.basename(filepath)
        # Ignorar temporales/ocultos/control
        if (
            filename in self.IGNORED_FILENAMES or
            filename.startswith("~$") or
            filename.startswith(".") or
            filename.lower().endswith(".pkl")
        ):
            return False
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.SUPPORTED_EXTENSIONS or ext in self.IMAGE_EXTENSIONS
    
    @staticmethod
    def calculate_content_hash(content: str) -> str:
        """
        Calcula hash MD5 del contenido
        
        Args:
            content: Contenido del documento
            
        Returns:
            Hash MD5 del contenido
        """
        return hashlib.md5(content.encode("utf-8")).hexdigest()
    
    def load_documents(self) -> List[Document]:
        """
        Carga todos los documentos e imágenes del directorio de conocimiento
        
        Returns:
            Lista de documentos cargados (incluyendo descripciones de imágenes)
        """
        documents = []
        
        if not os.path.exists(KNOWLEDGE_DIR):
            print(f"⚠️ Directorio de conocimiento no existe: {KNOWLEDGE_DIR}")
            return documents
            
        for filename in os.listdir(KNOWLEDGE_DIR):
            filepath = os.path.join(KNOWLEDGE_DIR, filename)
            
            # Solo procesar archivos, no directorios
            if not os.path.isfile(filepath):
                continue
                
            # Solo procesar archivos soportados
            if not self.is_supported_file(filepath):
                continue
            
            # Determinar tipo de archivo
            ext = os.path.splitext(filename)[1].lower()
            file_type = "imagen" if ext in self.IMAGE_EXTENSIONS else "documento"
            
            content = self.extract_text_from_file(filepath)
            
            if content and len(content.strip()) > 30:  # Filtrar contenido muy corto
                documents.append(Document(filename=filename, content=content))
                print(f"✅ {file_type.capitalize()} cargado: {filename}")
            else:
                print(f"⚠️ No se pudo cargar contenido suficiente de: {filename}")
        
        print(f"📚 Total de documentos e imágenes cargados: {len(documents)}")
        return documents
    
    @staticmethod
    def documents_to_tuples(documents: List[Document]) -> List[DocumentTuple]:
        """
        Convierte lista de documentos a lista de tuplas para compatibilidad
        
        Args:
            documents: Lista de objetos Document
            
        Returns:
            Lista de tuplas (filename, content)
        """
        return [doc.to_tuple() for doc in documents]